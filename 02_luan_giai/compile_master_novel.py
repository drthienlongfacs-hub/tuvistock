import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Mảng chứa các file thành phần
files = ['draft_novel_act1.md', 'draft_novel_act2.md', 'draft_novel_act3.md']
output_file = 'MASTER_THE_SWORD_LEAVES_THE_SHEATH.docx'

doc = Document()

# Thiết lập font mặc định cho tiểu thuyết
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13) # Cỡ chữ tiểu thuyết thường lớn hơn báo cáo

# Bìa lót / Title Page
doc.add_paragraph('\n\n\n\n')
title = doc.add_heading('THE SWORD LEAVES THE SHEATH', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_vi = doc.add_paragraph('THANH KIẾM RA KHỎI VỎ')
subtitle_vi.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_vi.runs[0].bold = True
subtitle_vi.runs[0].font.size = Pt(16)

chronicle = doc.add_paragraph('\nBiên niên ký năm 2026 của Long\n(Dựa trên Lá số Tử Vi: Nhâm Thân - Bính Ngọ)')
chronicle.alignment = WD_ALIGN_PARAGRAPH.CENTER
chronicle.runs[0].italic = True

doc.add_page_break()

# Đọc và định dạng nội dung
for i, file_name in enumerate(files):
    if not os.path.exists(file_name):
        print(f"Warning: File {file_name} not found.")
        continue
        
    with open(file_name, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Loại bỏ Tiêu đề lặp lại ở file 1 (để không bị lặp với bìa lót)
    if i == 0:
        content = re.sub(r'# THE SWORD LEAVES THE SHEATH.*?\n---', '', content, flags=re.DOTALL)

    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line or line == '---' or line == '(HẾT)':
            continue
            
        if line.startswith('## '):
            if i > 0 or not line.startswith('## LỜI MỞ ĐẦU'): 
                doc.add_page_break() # Sang trang mới cho mỗi Act
            h = doc.add_heading(line[3:].strip(), level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph() 
            
        elif line.startswith('### '):
            doc.add_paragraph()
            h2 = doc.add_heading(line[4:].strip(), level=2)
            h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            
        elif line.startswith('**Cảnh'):
            doc.add_paragraph() # Spacing
            p = doc.add_paragraph()
            r = p.add_run(line.strip('*'))
            r.bold = True
            
        elif line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            # Khúc trích dẫn nghiêng toàn bộ dòng
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line.strip('*'))
            r.italic = True
            
        else:
            # Đoạn văn tiểu thuyết chuẩn: Lùi đầu dòng, cách dòng 1.5
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                elif part.startswith('*') and part.endswith('*'):
                    p.add_run(part[1:-1]).italic = True
                else:
                    p.add_run(part)

doc.save(output_file)
print(f"Successfully compiled Master Novel to {output_file}")
