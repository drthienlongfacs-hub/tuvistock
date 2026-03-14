import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Mảng chứa các file thành phần
files = ['draft_book_part1.md', 'draft_book_part2.md', 'draft_book_part3.md']
output_file = 'THANH_KIEM_RA_KHOI_VO_2026.docx'

doc = Document()

# Thiết lập font mặc định
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Tiêu đề chính
title = doc.add_heading('THANH KIẾM RA KHỎI VỎ', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('Biên niên ký năm 2026 của Long')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].italic = True

for i, file_name in enumerate(files):
    if not os.path.exists(file_name):
        print(f"Warning: File {file_name} not found.")
        continue
        
    with open(file_name, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Loại bỏ Tittle/Subtitle trùng nếu đọc từ file 1
    if i == 0:
        content = re.sub(r'# THANH KIẾM RA KHỎI VỎ\n\*\*Biên niên ký năm 2026 của Long\*\*\n\n---', '', content)

    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line or line == '---' or line == '(HẾT)':
            # Chèn ngắt trang nếu gặp header lớn
            continue
            
        if line.startswith('## '):
            h = doc.add_heading(line[3:].strip(), level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph() # spacing
            
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
            
        elif line.startswith('**Cảnh'):
            # Xử lý format Cảnh phim
            p = doc.add_paragraph()
            r = p.add_run(line.strip('*'))
            r.bold = True
            
        elif line.startswith('*'):
            # Italic quotes or blocks
            txt = line.strip('*')
            p = doc.add_paragraph()
            r = p.add_run(txt)
            r.italic = True
            
        else:
            # Đoạn văn bình thường
            
            # Thay thế markdown bold **text** -> run bold
            parts = re.split(r'(\*\*.*?\*\*)', line)
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.line_spacing = 1.5
            
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                elif part.startswith('*') and part.endswith('*'):
                    p.add_run(part[1:-1]).italic = True
                else:
                    p.add_run(part)

doc.save(output_file)
print(f"Successfully compiled novel to {output_file}")
