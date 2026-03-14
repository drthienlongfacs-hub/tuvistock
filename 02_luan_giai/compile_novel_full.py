import os
import glob
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

BRAIN_DIR = "/Users/mac/.gemini/antigravity/brain/6fb82159-baa4-4e05-a903-c15a25927049"
CHAPTERS_DIR = "/Users/mac/Desktop/Downloads/Luan giai tu vi by BS Long/02_luan_giai"
OUTPUT_FILE = os.path.join(CHAPTERS_DIR, "THE_SWORD_LEAVES_THE_SHEATH_FULL_2026.docx")

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Find and add main cover
main_cover_pattern = os.path.join(BRAIN_DIR, "the_sword_leaves_the_sheath_cover_*.png")
main_covers = glob.glob(main_cover_pattern)
if main_covers:
    main_cover_path = sorted(main_covers)[-1]
    doc.add_picture(main_cover_path, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

# Main Title
title = doc.add_heading('THANH KIẾM RA KHỎI VỎ', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('Biên niên ký năm 2026 của Long')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].italic = True

doc.add_page_break()

# Loop through 12 chapters
for i in range(1, 13):
    chapter_num_str = f"{i:02d}"
    chapter_file = os.path.join(CHAPTERS_DIR, f"novel_chapter_{chapter_num_str}.md")
    
    if not os.path.exists(chapter_file):
        print(f"File not found: {chapter_file}")
        continue
        
    print(f"Processing Chapter {i}...")
    
    # Find and add chapter cover
    cover_pattern = os.path.join(BRAIN_DIR, f"chapter_{chapter_num_str}_cover_*.png")
    covers = glob.glob(cover_pattern)
    if covers:
        cover_path = sorted(covers)[-1]
        doc.add_picture(cover_path, width=Inches(5.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    with open(chapter_file, 'r', encoding='utf-8') as f:
        content = f.read()
        
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line or line == '---' or line == '(Hết Chương 1)' or line == '(Hết Truyện)' or line.startswith('(Hết Chương'):
            continue
            
        if line.startswith('# '):
            h = doc.add_heading(line[2:].strip(), level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            
        elif line.startswith('### '):
            h2 = doc.add_heading(line[4:].strip(), level=2)
            
        elif line.startswith('*(Tháng'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line.strip('*'))
            r.italic = True
            doc.add_paragraph()
            
        elif line.startswith('**['):
            # System text
            p = doc.add_paragraph()
            r = p.add_run(line.strip('*'))
            r.bold = True
            r.font.name = 'Courier New'
            
        elif line.startswith('*') and line.endswith('*'):
            # Italic full line
            p = doc.add_paragraph()
            r = p.add_run(line.strip('*'))
            r.italic = True
            
        elif line.startswith('- *') or line.startswith('- **'):
            p = doc.add_paragraph()
            p.add_run(line)
            p.paragraph_format.left_indent = Inches(0.5)
            
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.line_spacing = 1.3
            
            # Simple bold parsing
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    subparts = re.split(r'(\*.*?\*)', part)
                    for sp in subparts:
                        if sp.startswith('*') and sp.endswith('*') and len(sp)>2:
                            p.add_run(sp[1:-1]).italic = True
                        else:
                            p.add_run(sp)
                            
    if i < 12:
        doc.add_page_break()

doc.save(OUTPUT_FILE)
print(f"Successfully compiled full novel to {OUTPUT_FILE}")
