import sys
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Custom styles
    styles = doc.styles
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            p = doc.add_heading(level=0)
            p.add_run(line[2:].strip())
        elif line.startswith('## '):
            p = doc.add_heading(level=1)
            p.add_run(line[3:].strip())
        elif line.startswith('### '):
            p = doc.add_heading(level=2)
            p.add_run(line[4:].strip())
        elif line.startswith('#### '):
            p = doc.add_heading(level=3)
            p.add_run(line[5:].strip())
        elif line.startswith('> '):
            p = doc.add_paragraph()
            p.style = 'Quote'
            _add_formatted_text(p, line[2:].strip())
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, line[2:].strip())
        elif line == '---':
            doc.add_page_break()
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, line)
            
    doc.save(docx_path)
    print(f"Saved DOCX to {docx_path}")

def _add_formatted_text(p, text):
    # Extremely basic markdown bold parser
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            p.add_run(part[2:-2]).bold = True
        else:
            p.add_run(part)

if __name__ == '__main__':
    md_to_docx(sys.argv[1], sys.argv[2])
