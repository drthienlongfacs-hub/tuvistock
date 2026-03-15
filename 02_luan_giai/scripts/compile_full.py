import sys
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import docx

def create_docx(input_files, output_file):
    doc = Document()
    
    # Adjust normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    lines = []
    for f_in in input_files:
        with open(f_in, 'r', encoding='utf-8') as f:
            lines.extend(f.readlines())
            lines.append('\n')

    in_table = False
    table = None
    
    for line in lines:
        line_s = line.strip()
        if not line_s:
            if in_table:
                in_table = False
            continue
            
        if line_s.startswith('|') and line_s.endswith('|'):
            # Table processing
            if not in_table:
                in_table = True
                cols = len([x for x in line_s.split('|') if x.strip() or True]) - 2
                table = doc.add_table(rows=0, cols=cols)
                table.style = 'Table Grid'
            
            # if it's separator line
            if '---' in line_s:
                continue
            
            row_cells = table.add_row().cells
            cell_texts = [x.strip() for x in line_s.split('|')[1:-1]]
            for i, text in enumerate(cell_texts):
                if i < len(row_cells):
                    _add_formatted_text(row_cells[i].paragraphs[0], text)
            continue
        
        in_table = False
        
        if line_s.startswith('# '):
            p = doc.add_heading(level=0)
            p.add_run(line_s[2:].strip())
        elif line_s.startswith('## '):
            p = doc.add_heading(level=1)
            p.add_run(line_s[3:].strip())
        elif line_s.startswith('### '):
            p = doc.add_heading(level=2)
            p.add_run(line_s[4:].strip())
        elif line_s.startswith('#### '):
            p = doc.add_heading(level=3)
            p.add_run(line_s[5:].strip())
        elif line_s.startswith('> '):
            p = doc.add_paragraph()
            p.style = 'Quote' if 'Quote' in doc.styles else 'Normal'
            _add_formatted_text(p, line_s[2:].strip())
        elif line_s.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, line_s[2:].strip())
        elif line_s == '---' or line_s == '***':
            doc.add_page_break()
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, line_s)
            
    doc.save(output_file)
    print(f"Saved DOCX to {output_file}")

def _add_formatted_text(p, text):
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            p.add_run(part[2:-2]).bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            p.add_run(part[1:-1]).italic = True
        else:
            p.add_run(part)

if __name__ == '__main__':
    inputs = [
        "/Users/mac/Desktop/Downloads/Luan giai tu vi by BS Long/02_luan_giai/docx_phase_1.md",
        "/Users/mac/Desktop/Downloads/Luan giai tu vi by BS Long/02_luan_giai/docx_phase_2.md",
        "/Users/mac/Desktop/Downloads/Luan giai tu vi by BS Long/02_luan_giai/docx_phase_3.md",
        "/Users/mac/Desktop/Downloads/Luan giai tu vi by BS Long/02_luan_giai/docx_phase_4.md"
    ]
    create_docx(inputs, "/Users/mac/Desktop/Downloads/Luan giai tu vi by BS Long/02_luan_giai/luan_giai_toan_dien_long_2026.docx")
