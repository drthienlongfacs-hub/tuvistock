import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_dissertation_docx(input_files, output_file):
    doc = Document()
    
    # Configure default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.5
    
    # Configure heading styles
    for level in range(4):
        hs = doc.styles[f'Heading {level+1}'] if f'Heading {level+1}' in doc.styles else None
        if hs:
            hs.font.name = 'Times New Roman'
            if level == 0:
                hs.font.size = Pt(18)
                hs.font.bold = True
            elif level == 1:
                hs.font.size = Pt(16)
                hs.font.bold = True
            elif level == 2:
                hs.font.size = Pt(14)
                hs.font.bold = True
            elif level == 3:
                hs.font.size = Pt(13)
                hs.font.bold = True

    lines = []
    for f_in in input_files:
        with open(f_in, 'r', encoding='utf-8') as f:
            lines.extend(f.readlines())
            lines.append('\n')

    in_table = False
    table = None
    header_done = False
    
    for line in lines:
        line_s = line.strip()
        if not line_s:
            if in_table:
                in_table = False
                table = None
            continue
            
        # Table rows
        if line_s.startswith('|') and line_s.endswith('|'):
            cells = [x.strip() for x in line_s.split('|')[1:-1]]
            
            # Separator line
            if all(set(c) <= set('-: ') for c in cells):
                continue
                
            if not in_table:
                in_table = True
                ncols = len(cells)
                table = doc.add_table(rows=0, cols=ncols)
                table.style = 'Table Grid'
                # Header row
                header_done = False
            
            row = table.add_row()
            for i, text in enumerate(cells):
                if i < len(row.cells):
                    p = row.cells[i].paragraphs[0]
                    p.paragraph_format.space_after = Pt(2)
                    _add_formatted_text(p, text)
                    for run in p.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Times New Roman'
                    if not header_done:
                        for run in p.runs:
                            run.bold = True
            
            if not header_done:
                header_done = True
            continue
        
        in_table = False
        table = None
        
        # Headings
        if line_s.startswith('# ') and not line_s.startswith('## '):
            p = doc.add_heading(level=0)
            _add_formatted_text(p, line_s[2:].strip())
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line_s.startswith('## '):
            p = doc.add_heading(level=1)
            _add_formatted_text(p, line_s[3:].strip())
        elif line_s.startswith('### '):
            p = doc.add_heading(level=2)
            _add_formatted_text(p, line_s[4:].strip())
        elif line_s.startswith('#### '):
            p = doc.add_heading(level=3)
            _add_formatted_text(p, line_s[5:].strip())
        # Blockquote
        elif line_s.startswith('> '):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Cm(1.5)
            pf.space_before = Pt(4)
            pf.space_after = Pt(4)
            _add_formatted_text(p, line_s[2:].strip())
            for run in p.runs:
                run.italic = True
                run.font.size = Pt(12)
        # Lists
        elif line_s.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, line_s[2:].strip())
        # Horizontal rules / page breaks
        elif line_s == '---' or line_s == '***':
            doc.add_page_break()
        # Normal paragraph
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, line_s)
            
    doc.save(output_file)
    print(f"Saved dissertation DOCX to {output_file}")

def _add_formatted_text(p, text):
    # Handle bold (**text**) and italic (*text*)
    parts = re.split(r'(\*\*.*?\*\*|\*[^*]+?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Times New Roman'
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = p.add_run(part[1:-1])
            run.italic = True
            run.font.name = 'Times New Roman'
        else:
            run = p.add_run(part)
            run.font.name = 'Times New Roman'

if __name__ == '__main__':
    inputs = [
        "/Users/mac/Desktop/Downloads/Luan giai tu vi by BS Long/02_luan_giai/dissertation_batch1.md",
        "/Users/mac/Desktop/Downloads/Luan giai tu vi by BS Long/02_luan_giai/dissertation_batch2.md",
        "/Users/mac/Desktop/Downloads/Luan giai tu vi by BS Long/02_luan_giai/dissertation_batch3.md"
    ]
    create_dissertation_docx(inputs, "/Users/mac/Desktop/Downloads/Luan giai tu vi by BS Long/02_luan_giai/luan_giai_toan_dien_long_2026.docx")
